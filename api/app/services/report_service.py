"""Report service for generating DOCX audit reports."""

import io
import logging
from datetime import datetime
from typing import Any
from uuid import UUID

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession

from .case_service import case_service
from .embedding_service import embedding_service
from .ollama_service import ollama_service

logger = logging.getLogger(__name__)


class ReportService:
    """Service for generating DOCX audit reports."""

    # Severity colors for visual indicators
    SEVERITY_COLORS = {
        "CRITICAL": RGBColor(0xC0, 0x39, 0x2B),  # Red
        "HIGH": RGBColor(0xE7, 0x4C, 0x3C),  # Orange-red
        "MEDIUM": RGBColor(0xF3, 0x9C, 0x12),  # Orange
        "LOW": RGBColor(0x27, 0xAE, 0x60),  # Green
        "INFO": RGBColor(0x34, 0x98, 0xDB),  # Blue
    }

    # Status colors
    STATUS_COLORS = {
        "OPEN": RGBColor(0x34, 0x98, 0xDB),  # Blue
        "IN_PROGRESS": RGBColor(0xF3, 0x9C, 0x12),  # Orange
        "PENDING_REVIEW": RGBColor(0x95, 0x59, 0xE5),  # Purple
        "CLOSED": RGBColor(0x27, 0xAE, 0x60),  # Green
        "ARCHIVED": RGBColor(0x7F, 0x8C, 0x8D),  # Gray
    }

    def _setup_styles(self, doc: Document) -> None:
        """Set up custom styles for the document."""
        styles = doc.styles

        # Title style
        if "Report Title" not in [s.name for s in styles]:
            title_style = styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            title_style.paragraph_format.space_after = Pt(12)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Section heading style
        if "Section Heading" not in [s.name for s in styles]:
            section_style = styles.add_style("Section Heading", WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(6)

        # Subsection heading style
        if "Subsection Heading" not in [s.name for s in styles]:
            subsection_style = styles.add_style("Subsection Heading", WD_STYLE_TYPE.PARAGRAPH)
            subsection_style.font.size = Pt(12)
            subsection_style.font.bold = True
            subsection_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
            subsection_style.paragraph_format.space_before = Pt(12)
            subsection_style.paragraph_format.space_after = Pt(4)

    def _add_cover_page(
        self,
        doc: Document,
        case_data: dict[str, Any],
        title: str | None = None,
        watermark: str | None = None,
    ) -> None:
        """Add cover page to the report."""
        # Add some spacing at top
        for _ in range(3):
            doc.add_paragraph()

        # Report title
        title_text = title or f"Audit Case Report"
        title_para = doc.add_paragraph(title_text, style="Report Title")

        # Case ID prominently displayed
        case_id_para = doc.add_paragraph()
        case_id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_id_run = case_id_para.add_run(case_data.get("case_id", "Unknown"))
        case_id_run.font.size = Pt(24)
        case_id_run.font.bold = True
        case_id_run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)

        # Case title
        case_title_para = doc.add_paragraph()
        case_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_title_run = case_title_para.add_run(case_data.get("title", "Untitled Case"))
        case_title_run.font.size = Pt(16)
        case_title_run.font.italic = True

        doc.add_paragraph()

        # Classification/status box
        status = case_data.get("status", "OPEN")
        severity = case_data.get("severity", "MEDIUM")

        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Status: {status}  |  Severity: {severity}")

        # Watermark if provided
        if watermark:
            doc.add_paragraph()
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            watermark_run = watermark_para.add_run(watermark.upper())
            watermark_run.font.size = Pt(36)
            watermark_run.font.bold = True
            watermark_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Footer info
        for _ in range(5):
            doc.add_paragraph()

        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # Page break after cover
        doc.add_page_break()

    def _add_executive_summary(
        self,
        doc: Document,
        case_data: dict[str, Any],
        ai_summary: dict[str, Any] | None = None,
    ) -> None:
        """Add executive summary section."""
        doc.add_paragraph("Executive Summary", style="Section Heading")

        if ai_summary and ai_summary.get("summary"):
            # AI-generated summary
            summary_para = doc.add_paragraph(ai_summary["summary"])
            summary_para.paragraph_format.space_after = Pt(12)

            # Key points
            if ai_summary.get("key_points"):
                doc.add_paragraph("Key Points:", style="Subsection Heading")
                for point in ai_summary["key_points"]:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(point)

            # Risk assessment
            if ai_summary.get("risk_assessment"):
                doc.add_paragraph("Risk Assessment:", style="Subsection Heading")
                risk_para = doc.add_paragraph(ai_summary["risk_assessment"])

            # Recommended actions
            if ai_summary.get("recommended_actions"):
                doc.add_paragraph("Recommended Actions:", style="Subsection Heading")
                for action in ai_summary["recommended_actions"]:
                    bullet = doc.add_paragraph(style="List Bullet")
                    bullet.add_run(action)
        else:
            # Fallback to case summary if no AI summary
            summary = case_data.get("summary") or case_data.get("description") or "No summary available."
            doc.add_paragraph(summary)

        doc.add_paragraph()

    def _add_case_details(
        self,
        doc: Document,
        case_data: dict[str, Any],
        owner: dict[str, Any] | None = None,
    ) -> None:
        """Add case details section."""
        doc.add_paragraph("Case Details", style="Section Heading")

        # Create a table for case metadata
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        # Helper to add row
        def add_row(label: str, value: str) -> None:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else "N/A"
            # Bold the label
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        add_row("Case ID", case_data.get("case_id", "N/A"))
        add_row("Title", case_data.get("title", "N/A"))
        add_row("Scope", f"{case_data.get('scope_code', 'N/A')}")
        add_row("Type", case_data.get("case_type", "N/A"))
        add_row("Status", case_data.get("status", "N/A"))
        add_row("Severity", case_data.get("severity", "N/A"))

        if case_data.get("incident_date"):
            incident_date = case_data["incident_date"]
            if isinstance(incident_date, datetime):
                incident_date = incident_date.strftime("%Y-%m-%d")
            add_row("Incident Date", incident_date)

        if case_data.get("subject_user"):
            add_row("Subject User", case_data["subject_user"])

        if case_data.get("subject_computer"):
            add_row("Subject Computer", case_data["subject_computer"])

        if case_data.get("subject_devices"):
            devices = case_data["subject_devices"]
            if isinstance(devices, list):
                devices = ", ".join(devices)
            add_row("Subject Devices", devices)

        if case_data.get("related_users"):
            users = case_data["related_users"]
            if isinstance(users, list):
                users = ", ".join(users)
            add_row("Related Users", users)

        if owner:
            add_row("Case Owner", owner.get("full_name", owner.get("email", "N/A")))

        created_at = case_data.get("created_at")
        if created_at:
            if isinstance(created_at, datetime):
                created_at = created_at.strftime("%Y-%m-%d %H:%M")
            add_row("Created", created_at)

        if case_data.get("tags"):
            tags = case_data["tags"]
            if isinstance(tags, list):
                tags = ", ".join(tags)
            add_row("Tags", tags)

        doc.add_paragraph()

        # Description section
        if case_data.get("description"):
            doc.add_paragraph("Description:", style="Subsection Heading")
            doc.add_paragraph(case_data["description"])

        doc.add_paragraph()

    def _add_timeline(
        self,
        doc: Document,
        timeline_events: list[dict[str, Any]],
    ) -> None:
        """Add timeline section."""
        doc.add_paragraph("Timeline of Events", style="Section Heading")

        if not timeline_events:
            doc.add_paragraph("No timeline events recorded.")
            doc.add_paragraph()
            return

        # Create timeline table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        headers = ["Date/Time", "Event Type", "Description", "Source"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Add events
        for event in timeline_events:
            row = table.add_row()

            # Date/time
            event_time = event.get("event_time")
            if isinstance(event_time, datetime):
                event_time = event_time.strftime("%Y-%m-%d %H:%M")
            row.cells[0].text = str(event_time) if event_time else "N/A"

            # Event type
            row.cells[1].text = event.get("event_type", "N/A")

            # Description
            row.cells[2].text = event.get("description", "N/A")

            # Source
            row.cells[3].text = event.get("source", "N/A")

        doc.add_paragraph()

    def _add_findings(
        self,
        doc: Document,
        findings: list[dict[str, Any]],
    ) -> None:
        """Add findings section with severity indicators."""
        doc.add_paragraph("Findings", style="Section Heading")

        if not findings:
            doc.add_paragraph("No findings recorded.")
            doc.add_paragraph()
            return

        # Group by severity
        severity_order = ["CRITICAL", "HIGH", "MEDIUM", "LOW", "INFO"]
        findings_by_severity: dict[str, list] = {s: [] for s in severity_order}

        for finding in findings:
            severity = finding.get("severity", "MEDIUM")
            if severity in findings_by_severity:
                findings_by_severity[severity].append(finding)
            else:
                findings_by_severity["MEDIUM"].append(finding)

        finding_num = 1
        for severity in severity_order:
            severity_findings = findings_by_severity[severity]
            if not severity_findings:
                continue

            # Severity subsection
            severity_para = doc.add_paragraph(f"{severity} Severity Findings", style="Subsection Heading")

            for finding in severity_findings:
                # Finding title with number
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"Finding {finding_num}: {finding.get('title', 'Untitled')}")
                title_run.font.bold = True

                # Color code based on severity
                if severity in self.SEVERITY_COLORS:
                    title_run.font.color.rgb = self.SEVERITY_COLORS[severity]

                # Finding description
                if finding.get("description"):
                    doc.add_paragraph(finding["description"])

                finding_num += 1
                doc.add_paragraph()

    def _add_evidence_list(
        self,
        doc: Document,
        evidence: list[dict[str, Any]],
    ) -> None:
        """Add evidence list section with file hashes."""
        doc.add_paragraph("Evidence", style="Section Heading")

        if not evidence:
            doc.add_paragraph("No evidence files attached.")
            doc.add_paragraph()
            return

        # Evidence table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        header_cells = table.rows[0].cells
        headers = ["#", "File Name", "Type", "Size", "Hash (SHA-256)"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True

        # Add evidence items
        for idx, item in enumerate(evidence, 1):
            row = table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.get("file_name", "Unknown")
            row.cells[2].text = item.get("mime_type", "N/A")

            # Format file size
            size = item.get("file_size", 0)
            if size > 1024 * 1024:
                size_str = f"{size / (1024 * 1024):.2f} MB"
            elif size > 1024:
                size_str = f"{size / 1024:.2f} KB"
            else:
                size_str = f"{size} bytes"
            row.cells[3].text = size_str

            # Hash (truncated for display)
            file_hash = item.get("file_hash", "N/A")
            if file_hash and len(file_hash) > 16:
                file_hash = file_hash[:16] + "..."
            row.cells[4].text = file_hash

        doc.add_paragraph()

        # Evidence descriptions
        doc.add_paragraph("Evidence Descriptions:", style="Subsection Heading")
        for idx, item in enumerate(evidence, 1):
            if item.get("description"):
                item_para = doc.add_paragraph()
                item_para.add_run(f"{idx}. {item.get('file_name', 'Unknown')}: ").bold = True
                item_para.add_run(item["description"])

        doc.add_paragraph()

    def _add_similar_cases(
        self,
        doc: Document,
        similar_cases: list[dict[str, Any]],
    ) -> None:
        """Add similar cases appendix section."""
        doc.add_paragraph("Related Cases", style="Section Heading")

        if not similar_cases:
            doc.add_paragraph("No similar cases identified.")
            doc.add_paragraph()
            return

        for case in similar_cases:
            case_para = doc.add_paragraph()
            case_id = case.get("case_id", "Unknown")
            similarity = case.get("similarity_score", 0)

            case_para.add_run(f"{case_id}").bold = True
            case_para.add_run(f" (Similarity: {similarity:.0%})")

            if case.get("title"):
                doc.add_paragraph(f"  Title: {case['title']}")
            if case.get("status"):
                doc.add_paragraph(f"  Status: {case['status']}")

        doc.add_paragraph()

    def _add_entities_appendix(
        self,
        doc: Document,
        entities: list[dict[str, Any]],
    ) -> None:
        """Add extracted entities appendix."""
        doc.add_paragraph("Extracted Entities", style="Section Heading")

        if not entities:
            doc.add_paragraph("No entities extracted.")
            doc.add_paragraph()
            return

        # Group by type
        entities_by_type: dict[str, list] = {}
        for entity in entities:
            etype = entity.get("entity_type", "other")
            if etype not in entities_by_type:
                entities_by_type[etype] = []
            entities_by_type[etype].append(entity)

        for etype, type_entities in entities_by_type.items():
            doc.add_paragraph(f"{etype.replace('_', ' ').title()}:", style="Subsection Heading")

            for entity in type_entities:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(entity.get("value", "N/A"))

        doc.add_paragraph()

    async def generate_report(
        self,
        db: AsyncSession,
        case_id: str,
        template: str = "STANDARD",
        include_evidence: bool = True,
        include_similar: bool = True,
        include_ai_summary: bool = True,
        watermark: str | None = None,
        title: str | None = None,
    ) -> io.BytesIO:
        """
        Generate a DOCX report for a case.

        Args:
            db: Database session
            case_id: Case ID string (e.g., 'FIN-USB-0001')
            template: Report template (STANDARD, EXECUTIVE_SUMMARY, DETAILED, COMPLIANCE)
            include_evidence: Include evidence list section
            include_similar: Include similar cases section
            include_ai_summary: Include AI-generated summary
            watermark: Optional watermark text
            title: Optional custom title

        Returns:
            BytesIO buffer containing the DOCX file

        Raises:
            ValueError: If case not found
        """
        logger.info(f"Generating {template} report for case {case_id}")

        # Fetch case data
        case_data = await case_service.get_case(db, case_id)
        if not case_data:
            raise ValueError(f"Case not found: {case_id}")

        case_uuid = case_data["id"]

        # Fetch related data
        findings = await case_service.get_case_findings(db, case_uuid)
        timeline = await case_service.get_case_timeline(db, case_uuid)
        evidence = await case_service.get_case_evidence(db, case_uuid) if include_evidence else []

        # Get owner info
        owner = None
        if case_data.get("owner_id"):
            owner = await case_service.get_user_brief(db, case_data["owner_id"])

        # Get AI summary if requested
        ai_summary = None
        if include_ai_summary:
            try:
                ai_summary = await ollama_service.summarize_case_structured({
                    "case_id": case_data.get("case_id"),
                    "case_type": case_data.get("case_type"),
                    "severity": case_data.get("severity"),
                    "title": case_data.get("title"),
                    "description": case_data.get("description"),
                    "subject_user": case_data.get("subject_user"),
                    "subject_computer": case_data.get("subject_computer"),
                    "subject_devices": case_data.get("subject_devices"),
                    "findings": [
                        {"title": f.get("title"), "description": f.get("description"), "severity": f.get("severity")}
                        for f in findings
                    ],
                    "timeline_events": [
                        {"event_time": str(e.get("event_time")), "description": e.get("description")}
                        for e in timeline
                    ],
                })
            except Exception as e:
                logger.warning(f"Failed to generate AI summary: {e}")
                ai_summary = None

        # Get similar cases if requested
        similar_cases = []
        if include_similar:
            try:
                similar_cases = await embedding_service.find_similar_cases(
                    db, case_uuid, limit=5, min_similarity=0.6
                )
            except Exception as e:
                logger.warning(f"Failed to find similar cases: {e}")
                similar_cases = []

        # Get entities
        entities = []
        try:
            from sqlalchemy import text
            result = await db.execute(
                text("SELECT entity_type, value FROM case_entities WHERE case_id = :case_id ORDER BY entity_type, value"),
                {"case_id": str(case_uuid)}
            )
            entities = [dict(row._mapping) for row in result.fetchall()]
        except Exception as e:
            logger.warning(f"Failed to get entities: {e}")

        # Create document
        doc = Document()
        self._setup_styles(doc)

        # Generate sections based on template
        if template == "EXECUTIVE_SUMMARY":
            # Brief summary only
            self._add_cover_page(doc, case_data, title, watermark)
            self._add_executive_summary(doc, case_data, ai_summary)
            self._add_case_details(doc, case_data, owner)
            if findings:
                self._add_findings(doc, findings[:5])  # Top 5 findings only

        elif template == "DETAILED":
            # Full detailed report
            self._add_cover_page(doc, case_data, title, watermark)
            self._add_executive_summary(doc, case_data, ai_summary)
            self._add_case_details(doc, case_data, owner)
            self._add_timeline(doc, timeline)
            self._add_findings(doc, findings)
            self._add_evidence_list(doc, evidence)
            self._add_similar_cases(doc, similar_cases)
            self._add_entities_appendix(doc, entities)

        elif template == "COMPLIANCE":
            # Compliance-focused report
            self._add_cover_page(doc, case_data, title or "Compliance Report", watermark or "CONFIDENTIAL")
            self._add_case_details(doc, case_data, owner)
            self._add_findings(doc, findings)
            self._add_timeline(doc, timeline)
            self._add_evidence_list(doc, evidence)

        else:  # STANDARD
            # Standard report with all sections
            self._add_cover_page(doc, case_data, title, watermark)
            self._add_executive_summary(doc, case_data, ai_summary)
            self._add_case_details(doc, case_data, owner)
            self._add_timeline(doc, timeline)
            self._add_findings(doc, findings)
            if include_evidence:
                self._add_evidence_list(doc, evidence)
            if include_similar:
                self._add_similar_cases(doc, similar_cases)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info(f"Report generated successfully for case {case_id}")
        return buffer


# Singleton instance
report_service = ReportService()
